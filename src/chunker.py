from __future__ import annotations

import argparse
import json
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable


os.environ.setdefault("HF_HUB_OFFLINE", "1")
os.environ.setdefault("TRANSFORMERS_OFFLINE", "1")


PACKAGE_DIR = Path(__file__).resolve().parent
DEFAULT_MATERIAL_DIR = PACKAGE_DIR.parent / "data" / "material"
DEFAULT_TEXT_DIR = PACKAGE_DIR.parent / "data" / "text"
DEFAULT_CHUNKS_DIR = PACKAGE_DIR.parent / "data" / "chunks"
DEFAULT_DB_PATH = PACKAGE_DIR.parent / "data" / "chroma"
DEFAULT_COLLECTION_NAME = "thesis_chunks"
DEFAULT_MODEL_NAME = "BAAI/bge-small-zh-v1.5"
DEFAULT_LOCAL_MODEL_PATH = PACKAGE_DIR.parent / "models" / "bge-small-zh-v1.5"

SUPPORTED_EXTENSIONS = {".docx", ".txt", ".md", ".pdf"}

MIN_CHUNK_SIZE = 250
TARGET_CHUNK_SIZE = 850
MAX_CHUNK_SIZE = 1100
OVERLAP_PARAGRAPHS = 1

HEADING_PATTERNS = [
    r"^第[一二三四五六七八九十百千万\d]+[章节篇部分]\s*\S*",
    r"^\d+(?:\.\d+){0,3}\s+\S+",
    r"^\d+(?:[.．、]\d+)*[.．、]\s*\S+",
    r"^[一二三四五六七八九十]+[、.．]\s*\S+",
    r"^(中文摘要|英文摘要|摘要|Abstract|关键词|Keywords|目录|参考文献|致谢)$",
    r"^(毕业论文.*|毕业设计.*|指导教师审查意见|评阅教师评语|答辩会议记录)$",
]

SENTENCE_SPLIT_RE = re.compile(r"(?<=[。！？；.!?;])")
TOC_TITLE_RE = re.compile(r"^\s*目\s*录\s*$")
TOC_LINE_RE = re.compile(r"(\t+\s*\d+\s*$|\.{2,}\s*\d+\s*$|…{2,}\s*\d+\s*$)")


@dataclass
class SourceDocument:
    path: Path
    text: str


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="提取 Material 中的资料文本，按章节切割 chunk，并生成 Chroma 向量库。"
    )
    parser.add_argument(
        "--material-dir",
        type=Path,
        default=DEFAULT_MATERIAL_DIR,
        help="资料目录，默认是 ./data/material",
    )
    parser.add_argument(
        "--text-dir",
        type=Path,
        default=DEFAULT_TEXT_DIR,
        help="提取后的文本输出目录，默认是 ./data/text",
    )
    parser.add_argument(
        "--chunks-dir",
        type=Path,
        default=DEFAULT_CHUNKS_DIR,
        help="chunk JSON 输出目录，默认是 ./data/chunks",
    )
    parser.add_argument(
        "--db-path",
        type=Path,
        default=DEFAULT_DB_PATH,
        help="Chroma 向量库目录，默认是 ./data/chroma",
    )
    parser.add_argument(
        "--collection-name",
        default=DEFAULT_COLLECTION_NAME,
        help="Chroma collection 名称，默认是 thesis_chunks",
    )
    parser.add_argument(
        "--skip-embedding",
        action="store_true",
        help="只提取文本和切块，不生成向量库。",
    )
    return parser.parse_args()


def find_material_files(material_dir: Path) -> list[Path]:
    if not material_dir.exists():
        material_dir.mkdir(parents=True, exist_ok=True)

    files = []
    for path in material_dir.rglob("*"):
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            if path.name.startswith("~$"):
                continue
            files.append(path)

    return sorted(files, key=lambda item: item.name.lower())


def extract_docx(path: Path) -> str:
    from docx import Document

    document = Document(path)
    paragraphs = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)


def extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as error:
        raise RuntimeError("当前环境没有安装 pypdf，暂时无法提取 PDF。") from error

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages.append(text.strip())

    return "\n\n".join(pages)


def extract_plain_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(encoding="gb18030")


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return extract_docx(path)
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix in {".txt", ".md"}:
        return extract_plain_text(path)

    raise ValueError(f"不支持的文件类型：{path.suffix}")


def extract_sources(material_dir: Path) -> list[SourceDocument]:
    files = find_material_files(material_dir)
    if not files:
        raise FileNotFoundError(f"没有在资料目录中找到可处理文件：{material_dir}")

    sources = []
    skipped = []

    for path in files:
        try:
            text = extract_text(path).strip()
        except Exception as error:
            skipped.append((path, str(error)))
            continue

        if text:
            sources.append(SourceDocument(path=path, text=text))

    if skipped:
        print("以下文件未能提取，将跳过：")
        for path, reason in skipped:
            print(f"- {path.name}: {reason}")

    if not sources:
        raise RuntimeError("没有成功提取到任何文本，请检查 Material 目录中的资料文件。")

    return sources


def safe_stem(path: Path) -> str:
    stem = path.stem.strip() or "未命名文件"
    return re.sub(r'[<>:"/\\|?*\s]+', "_", stem).strip("_")


def output_name(prefix: str, source_path: Path, suffix: str) -> str:
    return f"{prefix}_{safe_stem(source_path)}{suffix}"


def write_text_outputs(sources: Iterable[SourceDocument], text_dir: Path) -> tuple[dict[str, Path], int]:
    text_dir.mkdir(parents=True, exist_ok=True)
    outputs = {}
    total_chars = 0

    for source in sources:
        text = source.text.strip()
        output_path = text_dir / output_name("thesis", source.path, ".txt")
        output_path.write_text(text, encoding="utf-8")
        outputs[source.path.name] = output_path
        total_chars += len(text)

    return outputs, total_chars


def split_paragraphs(text: str) -> list[str]:
    paragraphs = []
    for line in text.splitlines():
        line = line.strip()
        if line:
            paragraphs.append(line)
    return paragraphs


def remove_table_of_contents(paragraphs: list[str]) -> list[str]:
    cleaned = []
    for paragraph in paragraphs:
        if TOC_TITLE_RE.match(paragraph):
            continue
        if TOC_LINE_RE.search(paragraph):
            continue
        cleaned.append(paragraph)
    return cleaned


def is_heading(paragraph: str) -> bool:
    if "\t" in paragraph:
        return False
    if len(paragraph) > 40:
        return False
    if re.match(r"^\d+[.．、]\S+[：:]", paragraph):
        return False
    return any(re.match(pattern, paragraph) for pattern in HEADING_PATTERNS)


def inline_section_title(paragraph: str) -> str | None:
    compact = re.sub(r"\s+", "", paragraph)
    if compact.startswith("[摘要]") or compact.startswith("摘要："):
        return "中文摘要"
    if compact.startswith("[关键词]") or compact.startswith("关键词："):
        return "关键词"
    if compact.startswith("[Abstract]") or compact.startswith("Abstract:"):
        return "英文摘要"
    if compact.startswith("[Keywords]") or compact.startswith("Keywords:"):
        return "Keywords"
    return None


def split_into_sections(source: SourceDocument) -> list[dict]:
    raw_paragraphs = split_paragraphs(source.text)
    paragraphs = remove_table_of_contents(raw_paragraphs)
    sections = []
    current_title = "文档前置信息"
    current_paragraphs = []

    for paragraph in paragraphs:
        inline_title = inline_section_title(paragraph)

        if inline_title:
            if current_paragraphs:
                sections.append(
                    {
                        "source": source.path.name,
                        "section_title": current_title,
                        "paragraphs": current_paragraphs,
                    }
                )
            current_title = inline_title
            current_paragraphs = [paragraph]
            continue

        if is_heading(paragraph):
            if current_paragraphs:
                sections.append(
                    {
                        "source": source.path.name,
                        "section_title": current_title,
                        "paragraphs": current_paragraphs,
                    }
                )
            current_title = paragraph
            current_paragraphs = [paragraph]
            continue

        current_paragraphs.append(paragraph)

    if current_paragraphs:
        sections.append(
            {
                "source": source.path.name,
                "section_title": current_title,
                "paragraphs": current_paragraphs,
            }
        )

    return sections


def split_long_paragraph(paragraph: str, max_size: int) -> list[str]:
    sentences = [part.strip() for part in SENTENCE_SPLIT_RE.split(paragraph) if part.strip()]
    if not sentences:
        return [paragraph]

    parts = []
    current = ""

    for sentence in sentences:
        if len(sentence) > max_size:
            if current:
                parts.append(current)
                current = ""
            for start in range(0, len(sentence), max_size):
                parts.append(sentence[start : start + max_size].strip())
            continue

        candidate = current + sentence if current else sentence
        if len(candidate) <= max_size:
            current = candidate
        else:
            if current:
                parts.append(current)
            current = sentence

    if current:
        parts.append(current)

    return parts


def flush_chunk(chunks: list[dict], source: str, section_title: str, paragraphs: list[str]) -> None:
    text = "\n".join(paragraphs).strip()
    if not text:
        return

    chunks.append(
        {
            "source": source,
            "section_title": section_title,
            "text": text,
            "char_count": len(text),
        }
    )


def build_section_chunks(section: dict) -> list[dict]:
    source = section["source"]
    section_title = section["section_title"]
    paragraphs = []

    for paragraph in section["paragraphs"]:
        if len(paragraph) <= MAX_CHUNK_SIZE:
            paragraphs.append(paragraph)
        else:
            paragraphs.extend(split_long_paragraph(paragraph, TARGET_CHUNK_SIZE))

    chunks = []
    current = []

    for paragraph in paragraphs:
        candidate = "\n".join(current + [paragraph]).strip()

        if current and len(candidate) > MAX_CHUNK_SIZE:
            flush_chunk(chunks, source, section_title, current)
            overlap = current[-OVERLAP_PARAGRAPHS:] if len(current) > 1 else []
            current = overlap + [paragraph]
            continue

        current.append(paragraph)

        if len("\n".join(current)) >= TARGET_CHUNK_SIZE:
            flush_chunk(chunks, source, section_title, current)
            current = current[-OVERLAP_PARAGRAPHS:] if len(current) > 1 else []

    if current:
        tail_text = "\n".join(current).strip()
        if chunks and chunks[-1]["text"].endswith(tail_text):
            return chunks

        if (
            chunks
            and len(tail_text) < MIN_CHUNK_SIZE
            and len(chunks[-1]["text"]) + len(tail_text) + 1 <= MAX_CHUNK_SIZE
        ):
            chunks[-1]["text"] = chunks[-1]["text"] + "\n" + tail_text
            chunks[-1]["char_count"] = len(chunks[-1]["text"])
        else:
            flush_chunk(chunks, source, section_title, current)

    return chunks


def can_merge_chunks(left: dict, right: dict) -> bool:
    if left["source"] != right["source"]:
        return False
    if left["section_title"] == right["section_title"]:
        return True

    hard_boundaries = [
        "摘要",
        "关键词",
        "Abstract",
        "Keywords",
        "参考文献",
        "致谢",
        "所需资料",
        "主要参考文献",
    ]

    left_title = left["section_title"]
    right_title = right["section_title"]

    if any(boundary in left_title for boundary in hard_boundaries):
        return False
    if any(boundary in right_title for boundary in hard_boundaries):
        return False

    return len(left["text"]) < MIN_CHUNK_SIZE


def merge_small_chunks(chunks: list[dict]) -> list[dict]:
    merged = []
    index = 0

    while index < len(chunks):
        chunk = chunks[index].copy()

        while (
            len(chunk["text"]) < MIN_CHUNK_SIZE
            and index + 1 < len(chunks)
            and len(chunk["text"]) + len(chunks[index + 1]["text"]) + 1 <= MAX_CHUNK_SIZE
            and can_merge_chunks(chunk, chunks[index + 1])
        ):
            index += 1
            next_chunk = chunks[index]
            chunk["text"] = chunk["text"] + "\n" + next_chunk["text"]
            chunk["char_count"] = len(chunk["text"])
            if next_chunk["section_title"] not in chunk["section_title"]:
                chunk["section_title"] = chunk["section_title"] + " / " + next_chunk["section_title"]

        if (
            len(chunk["text"]) < MIN_CHUNK_SIZE
            and merged
            and len(merged[-1]["text"]) + len(chunk["text"]) + 1 <= MAX_CHUNK_SIZE
            and can_merge_chunks(merged[-1], chunk)
        ):
            merged[-1]["text"] = merged[-1]["text"] + "\n" + chunk["text"]
            merged[-1]["char_count"] = len(merged[-1]["text"])
            if chunk["section_title"] not in merged[-1]["section_title"]:
                merged[-1]["section_title"] = merged[-1]["section_title"] + " / " + chunk["section_title"]
        else:
            merged.append(chunk)

        index += 1

    return merged


def build_chunks(sources: Iterable[SourceDocument]) -> tuple[list[dict], int]:
    sections = []
    for source in sources:
        sections.extend(split_into_sections(source))

    chunks = []
    for section in sections:
        chunks.extend(build_section_chunks(section))

    chunks = merge_small_chunks(chunks)

    for index, chunk in enumerate(chunks, start=1):
        chunk["chunk_id"] = f"chunk_{index:03d}"

    return chunks, len(sections)


def write_chunk_outputs(chunks: list[dict], sources: Iterable[SourceDocument], chunks_dir: Path) -> dict[str, Path]:
    chunks_dir.mkdir(parents=True, exist_ok=True)
    outputs = {}

    chunks_by_source = {}
    for chunk in chunks:
        chunks_by_source.setdefault(chunk.get("source", ""), []).append(chunk)

    for source in sources:
        source_chunks = chunks_by_source.get(source.path.name, [])
        output_path = chunks_dir / output_name("chunks", source.path, ".json")
        output_path.write_text(
            json.dumps(source_chunks, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        outputs[source.path.name] = output_path

    return outputs


def load_embedding_model(local_model_path: Path, model_name: str = DEFAULT_MODEL_NAME):
    from sentence_transformers import SentenceTransformer

    if local_model_path.exists():
        return SentenceTransformer(str(local_model_path))

    try:
        return SentenceTransformer(model_name, local_files_only=True)
    except TypeError:
        return SentenceTransformer(model_name)
    except Exception as error:
        raise RuntimeError(
            "没有找到可用的本地 embedding 模型缓存。请把 BAAI/bge-small-zh-v1.5 "
            "放到 models/bge-small-zh-v1.5，或在联网环境下先下载模型。"
        ) from error


def build_vector_db(
    chunks: list[dict],
    db_path: Path,
    collection_name: str,
    local_model_path: Path = DEFAULT_LOCAL_MODEL_PATH,
) -> int:
    import chromadb

    print("正在加载 embedding 模型...")
    model = load_embedding_model(local_model_path)

    texts = [chunk["text"] for chunk in chunks]
    ids = [chunk["chunk_id"] for chunk in chunks]
    metadatas = [
        {
            "chunk_id": chunk["chunk_id"],
            "source": chunk.get("source", ""),
            "section_title": chunk.get("section_title", ""),
            "char_count": chunk.get("char_count", 0),
        }
        for chunk in chunks
    ]

    print("正在生成 embedding...")
    embeddings = model.encode(texts, normalize_embeddings=True).tolist()
    embedding_dim = len(embeddings[0]) if embeddings else 0

    print("正在写入 Chroma 向量库...")
    client = chromadb.PersistentClient(path=str(db_path))
    try:
        client.delete_collection(collection_name)
    except Exception:
        pass

    collection = client.get_or_create_collection(name=collection_name)
    collection.add(
        ids=ids,
        documents=texts,
        metadatas=metadatas,
        embeddings=embeddings,
    )

    return embedding_dim


def print_preview(chunks: list[dict]) -> None:
    print()
    print("前 3 个 chunk 预览：")
    for chunk in chunks[:3]:
        print("-" * 50)
        print(chunk["chunk_id"])
        print(f"来源文件：{chunk.get('source', '')}")
        print(f"所属章节：{chunk.get('section_title', '')}")
        print(f"字符数：{chunk.get('char_count', 0)}")
        print(chunk["text"][:300])


def main() -> None:
    args = parse_args()

    print(f"资料目录：{args.material_dir}")
    sources = extract_sources(args.material_dir)
    text_outputs, total_chars = write_text_outputs(sources, args.text_dir)
    chunks, section_count = build_chunks(sources)
    chunk_outputs = write_chunk_outputs(chunks, sources, args.chunks_dir)

    raw_paragraph_count = sum(len(split_paragraphs(source.text)) for source in sources)
    chunk_count = len(chunks)
    avg_chars = sum(chunk["char_count"] for chunk in chunks) / chunk_count if chunk_count else 0

    print("文本提取与章节切割完成。")
    print(f"处理文件数：{len(sources)}")
    print(f"提取文本输出目录：{args.text_dir}")
    print(f"chunk 输出目录：{args.chunks_dir}")
    for source in sources:
        print(f"- {source.path.name}")
        print(f"  文本：{text_outputs[source.path.name]}")
        print(f"  chunks：{chunk_outputs[source.path.name]}")
    print(f"提取文本总字符数：{total_chars}")
    print(f"原始段落数量：{raw_paragraph_count}")
    print(f"识别章节数：{section_count}")
    print(f"chunk 数量：{chunk_count}")
    print(f"平均 chunk 字符数：{avg_chars:.2f}")

    if not args.skip_embedding:
        embedding_dim = build_vector_db(
            chunks=chunks,
            db_path=args.db_path,
            collection_name=args.collection_name,
        )
        print("向量库构建完成。")
        print(f"写入 collection：{args.collection_name}")
        print(f"向量库路径：{args.db_path}")
        print(f"embedding 维度：{embedding_dim}")

    print_preview(chunks)


if __name__ == "__main__":
    main()
